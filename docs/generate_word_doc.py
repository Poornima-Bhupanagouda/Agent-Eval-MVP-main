#!/usr/bin/env python3
"""Generate Word document for Lilly Agent Eval Developer Guide."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('Lilly Agent Eval', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Developer & Architecture Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)

    version = doc.add_paragraph('Version 3.0.0 | March 2026')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. What is This Product?',
        '2. The Big Picture (Architecture)',
        '3. Project Files Explained',
        '4. Core Components (The Engine)',
        '5. Data Models (How We Store Information)',
        '6. API Endpoints (How Frontend Talks to Backend)',
        '7. The User Interface',
        '8. Database Tables',
        '9. Getting Started',
        '10. How to Add New Features',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ===========================================
    # Section 1: What is This Product?
    # ===========================================
    doc.add_heading('1. What is This Product?', level=1)

    doc.add_heading('The Problem We Solve', level=2)
    doc.add_paragraph(
        'When companies build AI chatbots (we call them "agents"), they need to answer important questions:'
    )
    problems = [
        'Does my chatbot give correct answers?',
        'Is my chatbot safe? Does it ever say harmful things?',
        'Which version of my chatbot is better - the old one or the new one?',
        'If I chain multiple chatbots together, does the whole pipeline work?',
        'How has my chatbot performed over the last 3 months?',
    ]
    for p in problems:
        doc.add_paragraph(f'• {p}')

    doc.add_paragraph('')
    doc.add_paragraph(
        'Lilly Agent Eval answers all these questions. It is a testing tool for AI chatbots.'
    ).runs[0].bold = True

    doc.add_heading('Who Uses This?', level=2)
    users = [
        ('AI Developers', 'To test if their chatbot works correctly'),
        ('QA Teams', 'To run automated tests before releasing new versions'),
        ('Product Managers', 'To compare different chatbot versions'),
        ('Data Scientists', 'To measure chatbot quality with metrics'),
    ]
    for user, reason in users:
        doc.add_paragraph(f'{user}: {reason}', style='List Bullet')

    doc.add_heading('Key Features', level=2)
    features = [
        ('Quick Test', 'Test any chatbot instantly with one click'),
        ('Agent Registry', 'Keep a list of all your chatbots in one place'),
        ('A/B Testing', 'Scientifically compare two chatbots to find the better one'),
        ('Chain Testing', 'Test chatbots that work together in sequence'),
        ('Test Suites', 'Save groups of tests to run again and again'),
        ('Evaluation Metrics', 'Measure quality: Is the answer relevant? Is it safe? Is it accurate?'),
        ('History & Analytics', 'See how your chatbot performed over time'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Feature'
    hdr[1].text = 'Why You Need It'
    for feat, why in features:
        row = table.add_row().cells
        row[0].text = feat
        row[1].text = why

    doc.add_page_break()

    # ===========================================
    # Section 2: Architecture
    # ===========================================
    doc.add_heading('2. The Big Picture (Architecture)', level=1)

    doc.add_heading('Why Do We Need Architecture?', level=2)
    doc.add_paragraph(
        'Architecture is like a building blueprint. It shows how all the parts connect. '
        'Without good architecture, the code becomes messy and hard to change. '
        'Our architecture is simple on purpose - it makes the system easy to understand and modify.'
    )

    doc.add_heading('The Three Layers', level=2)
    doc.add_paragraph(
        'We split the system into 3 layers. Each layer has one job. This separation makes it easy '
        'to change one part without breaking others.'
    )

    layers = [
        ('Layer 1: User Interface (Browser)',
         'This is what users see and click. It runs in a web browser.',
         'Users need a simple way to interact with the system without writing code.'),

        ('Layer 2: Backend Server (FastAPI)',
         'This is the brain. It receives requests, processes them, and sends responses.',
         'We need a central place to handle all the logic - calling chatbots, running evaluations, storing results.'),

        ('Layer 3: Database (SQLite)',
         'This is the memory. It stores all tests, results, and settings.',
         'We need to remember everything - past tests, registered chatbots, comparison results - so users can see history.'),
    ]

    for name, what, why in layers:
        doc.add_heading(name, level=3)
        doc.add_paragraph(f'What it is: {what}')
        doc.add_paragraph(f'Why we need it: {why}').runs[0].italic = True

    doc.add_heading('How They Talk to Each Other', level=2)
    flow = '''
User clicks "Run Test" in browser
        ↓
Browser sends request to Backend Server
        ↓
Backend calls the AI Chatbot being tested
        ↓
Backend evaluates the chatbot's response
        ↓
Backend saves results to Database
        ↓
Backend sends results back to Browser
        ↓
User sees pass/fail and scores
'''
    para = doc.add_paragraph()
    run = para.add_run(flow)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_page_break()

    # ===========================================
    # Section 3: Project Structure
    # ===========================================
    doc.add_heading('3. Project Files Explained', level=1)

    doc.add_heading('Why This Structure?', level=2)
    doc.add_paragraph(
        'We organize files by what they do. All the "core logic" lives in one folder. '
        'All the "web stuff" lives in another folder. This makes it easy to find things.'
    )

    doc.add_heading('Folder Map', level=2)

    structure_items = [
        ('agent_eval/', 'Main package folder - everything lives inside here'),
        ('  __init__.py', 'Makes this folder a Python package (required by Python)'),
        ('  cli.py', 'Command line tool (if users prefer terminal over browser)'),
        ('  core/', 'THE HEART OF THE SYSTEM - all business logic'),
        ('    models.py', 'Defines what a "Test", "Suite", "Result" looks like'),
        ('    executor.py', 'Calls AI chatbots over HTTP'),
        ('    evaluator.py', 'Measures if chatbot responses are good or bad'),
        ('    storage.py', 'Saves and retrieves data from database'),
        ('    statistics.py', 'Math for A/B testing (determines which chatbot wins)'),
        ('    file_parser.py', 'Reads PDF, Word, Excel files'),
        ('    introspector.py', 'Asks a chatbot "what can you do?"'),
        ('    context_generator.py', 'Creates sample test data'),
        ('    report_generator.py', 'Makes pretty HTML reports'),
        ('  web/', 'EVERYTHING FOR THE WEB INTERFACE'),
        ('    app.py', 'THE MAIN FILE - FastAPI server, all API routes'),
        ('    templates/', 'HTML files'),
        ('      index.html', 'THE USER INTERFACE - single page web app'),
    ]

    for item, desc in structure_items:
        p = doc.add_paragraph()
        run1 = p.add_run(item + '  ')
        run1.font.name = 'Courier New'
        run1.font.size = Pt(9)
        run2 = p.add_run('→ ' + desc)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    # ===========================================
    # Section 4: Core Components
    # ===========================================
    doc.add_heading('4. Core Components (The Engine)', level=1)

    doc.add_paragraph(
        'These are the parts that do the real work. Each component has ONE specific job. '
        'This is called "separation of concerns" - it keeps code clean and testable.'
    ).runs[0].italic = True

    # Executor
    doc.add_heading('4.1 Executor (executor.py)', level=2)
    doc.add_paragraph('ONE SENTENCE: Calls AI chatbots and gets their responses.').runs[0].bold = True

    doc.add_paragraph('WHY WE NEED IT:')
    doc.add_paragraph(
        'AI chatbots live on the internet as web services. To test them, we need to send HTTP requests '
        'and read the responses. The Executor handles all this network communication.'
    )

    doc.add_paragraph('THE PROBLEM IT SOLVES:')
    doc.add_paragraph(
        'Different chatbots expect different formats. Some want {"message": "hello"}, '
        'others want {"input": "hello"}, others want OpenAI format. The Executor tries '
        'multiple formats automatically so users don\'t have to figure it out.'
    )

    doc.add_paragraph('WHAT IT DOES:')
    executor_jobs = [
        'Sends HTTP POST requests to chatbot URLs',
        'Tries different request formats until one works',
        'Extracts the response text from various response formats',
        'Measures how long the chatbot took to respond (latency)',
        'Handles authentication (API keys, tokens, etc.)',
    ]
    for job in executor_jobs:
        doc.add_paragraph(job, style='List Bullet')

    # Evaluator
    doc.add_heading('4.2 Evaluator (evaluator.py)', level=2)
    doc.add_paragraph('ONE SENTENCE: Judges if a chatbot response is good or bad.').runs[0].bold = True

    doc.add_paragraph('WHY WE NEED IT:')
    doc.add_paragraph(
        'Getting a response from a chatbot is not enough. We need to know: Is this response correct? '
        'Is it safe? Is it relevant? The Evaluator answers these questions by running quality checks.'
    )

    doc.add_paragraph('THE PROBLEM IT SOLVES:')
    doc.add_paragraph(
        'Manually reading chatbot responses is slow and inconsistent. Different people might judge '
        'the same response differently. The Evaluator provides consistent, automated quality measurement.'
    )

    doc.add_paragraph('THE 7 QUALITY METRICS:')
    metrics = [
        ('Answer Relevancy', 'Does the answer actually address the question asked?'),
        ('Toxicity', 'Is the response safe? No harmful, offensive, or dangerous content?'),
        ('Bias', 'Is the response fair? Does it avoid stereotypes and discrimination?'),
        ('Faithfulness', 'Is the answer based on the provided documents (for RAG chatbots)?'),
        ('Hallucination', 'Did the chatbot make up facts that aren\'t true?'),
        ('Context Relevancy', 'Are the retrieved documents actually relevant to the question?'),
        ('Similarity', 'How close is the response to the expected answer?'),
    ]

    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = 'Table Grid'
    hdr = metrics_table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'What Question It Answers'
    for name, question in metrics:
        row = metrics_table.add_row().cells
        row[0].text = name
        row[1].text = question

    doc.add_paragraph('')
    doc.add_paragraph('TWO MODES:')
    doc.add_paragraph(
        '• DeepEval Mode: Uses OpenAI to intelligently judge responses (more accurate, requires API key)'
    )
    doc.add_paragraph(
        '• Heuristic Mode: Uses keyword matching and patterns (less accurate, but works offline)'
    )

    # Storage
    doc.add_heading('4.3 Storage (storage.py)', level=2)
    doc.add_paragraph('ONE SENTENCE: Saves everything to a database file.').runs[0].bold = True

    doc.add_paragraph('WHY WE NEED IT:')
    doc.add_paragraph(
        'Users want to see their test history. They want to compare results over time. '
        'They want their registered chatbots to still be there when they restart the app. '
        'Storage makes everything persistent.'
    )

    doc.add_paragraph('THE PROBLEM IT SOLVES:')
    doc.add_paragraph(
        'Without storage, all data disappears when you close the app. Storage saves everything '
        'to a SQLite file so nothing is lost. SQLite is a simple database that needs no setup - '
        'it\'s just one file on your computer.'
    )

    doc.add_paragraph('WHAT IT STORES:')
    storage_items = [
        'Test Suites - Collections of tests you created',
        'Test Cases - Individual test inputs and expected outputs',
        'Results - Every evaluation result with scores and metrics',
        'Agents - Registered chatbots with their URLs and settings',
        'A/B Comparisons - Results of chatbot vs chatbot tests',
        'Chains - Chatbot pipelines you configured',
    ]
    for item in storage_items:
        doc.add_paragraph(item, style='List Bullet')

    # Statistics
    doc.add_heading('4.4 Statistics (statistics.py)', level=2)
    doc.add_paragraph('ONE SENTENCE: Does the math to determine which chatbot is better.').runs[0].bold = True

    doc.add_paragraph('WHY WE NEED IT:')
    doc.add_paragraph(
        'When comparing two chatbots, you can\'t just say "this one scored 85% and that one scored 83%, '
        'so the first one wins." That 2% difference might be random luck. Statistics tells us if the '
        'difference is real or just noise.'
    )

    doc.add_paragraph('THE PROBLEM IT SOLVES:')
    doc.add_paragraph(
        'Without statistics, you might choose the wrong chatbot based on random variation. '
        'This module uses Welch\'s t-test (a standard scientific method) to calculate if one '
        'chatbot is truly better than another with statistical confidence.'
    )

    doc.add_paragraph('WHAT IT TELLS YOU:')
    stats_output = [
        'p-value: Probability the results are due to chance (lower = more confident)',
        'Effect size: How big is the difference (negligible/small/medium/large)',
        'Winner: Which chatbot is better, or if it\'s a tie',
    ]
    for item in stats_output:
        doc.add_paragraph(item, style='List Bullet')

    # File Parser
    doc.add_heading('4.5 File Parser (file_parser.py)', level=2)
    doc.add_paragraph('ONE SENTENCE: Reads documents so chatbots can use them as context.').runs[0].bold = True

    doc.add_paragraph('WHY WE NEED IT:')
    doc.add_paragraph(
        'Many chatbots are "RAG" chatbots - they answer questions based on documents you give them. '
        'Users want to upload a PDF of company policies and test if the chatbot can answer questions '
        'about it. The File Parser extracts text from various file formats.'
    )

    doc.add_paragraph('SUPPORTED FORMATS:')
    doc.add_paragraph('PDF, Word (.docx), Excel (.xlsx), CSV, Markdown, Text files')

    # Introspector
    doc.add_heading('4.6 Introspector (introspector.py)', level=2)
    doc.add_paragraph('ONE SENTENCE: Asks a chatbot what it can do.').runs[0].bold = True

    doc.add_paragraph('WHY WE NEED IT:')
    doc.add_paragraph(
        'When you register a new chatbot, you might not know its capabilities. Is it an HR chatbot? '
        'A customer support bot? Does it use documents? The Introspector asks the chatbot directly '
        'and figures out what type it is.'
    )

    # Report Generator
    doc.add_heading('4.7 Report Generator (report_generator.py)', level=2)
    doc.add_paragraph('ONE SENTENCE: Creates downloadable HTML reports.').runs[0].bold = True

    doc.add_paragraph('WHY WE NEED IT:')
    doc.add_paragraph(
        'Users want to share test results with their team or management. They need a nice-looking '
        'report they can email or present. The Report Generator creates professional HTML reports '
        'with Lilly branding.'
    )

    doc.add_page_break()

    # ===========================================
    # Section 5: Data Models
    # ===========================================
    doc.add_heading('5. Data Models (How We Store Information)', level=1)

    doc.add_paragraph('WHY WE NEED DATA MODELS:')
    doc.add_paragraph(
        'Data models define the shape of our data. They are like forms with specific fields. '
        'For example, a "Test" always has an input field. A "Result" always has a score. '
        'This consistency makes the code predictable and prevents errors.'
    )

    doc.add_heading('The Core Models', level=2)

    # Test
    doc.add_heading('Test', level=3)
    doc.add_paragraph('WHAT IT REPRESENTS: A single test case to run against a chatbot.')
    doc.add_paragraph('WHY IT EXISTS: We need a standard way to define "what question to ask" and "what answer we expect".')
    test_fields = '''
Test
├── input         → The question to send to the chatbot
├── expected      → The answer we expect (optional)
├── context       → Documents the chatbot should use (optional)
└── metrics       → Which quality checks to run (optional)
'''
    para = doc.add_paragraph()
    run = para.add_run(test_fields)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Suite
    doc.add_heading('Suite', level=3)
    doc.add_paragraph('WHAT IT REPRESENTS: A collection of related tests.')
    doc.add_paragraph('WHY IT EXISTS: Users want to group tests together. For example, "HR Policy Tests" might contain 20 questions about PTO, benefits, etc.')
    suite_fields = '''
Suite
├── name          → Name of the test collection
├── description   → What these tests are for
└── tests         → List of Test objects in this suite
'''
    para = doc.add_paragraph()
    run = para.add_run(suite_fields)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Result
    doc.add_heading('Result', level=3)
    doc.add_paragraph('WHAT IT REPRESENTS: The outcome of running one test.')
    doc.add_paragraph('WHY IT EXISTS: We need to record what happened - what the chatbot said, how it scored, how fast it was.')
    result_fields = '''
Result
├── input         → The question that was asked
├── output        → The chatbot's actual response
├── score         → Overall quality score (0-100)
├── passed        → Did it pass all quality checks?
├── latency_ms    → How long the chatbot took to respond
└── evaluations   → Individual scores for each metric
'''
    para = doc.add_paragraph()
    run = para.add_run(result_fields)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # RegisteredAgent
    doc.add_heading('RegisteredAgent', level=3)
    doc.add_paragraph('WHAT IT REPRESENTS: A chatbot that has been registered in the system.')
    doc.add_paragraph('WHY IT EXISTS: Users don\'t want to type the chatbot URL every time. They register it once and refer to it by name.')
    agent_fields = '''
RegisteredAgent
├── name          → Display name (e.g., "HR Bot v2.1")
├── endpoint      → URL where the chatbot lives
├── agent_type    → Type: "rag", "conversational", "simple"
├── auth_type     → How to authenticate: "none", "bearer_token", "api_key"
├── auth_config   → Authentication credentials
├── version       → Version tag for tracking different versions
└── is_active     → Is this chatbot currently available?
'''
    para = doc.add_paragraph()
    run = para.add_run(agent_fields)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # ABComparison
    doc.add_heading('ABComparison', level=3)
    doc.add_paragraph('WHAT IT REPRESENTS: A comparison between two chatbots.')
    doc.add_paragraph('WHY IT EXISTS: To track A/B test results - which chatbot won and by how much.')
    ab_fields = '''
ABComparison
├── agent_a_id    → First chatbot (the "control")
├── agent_b_id    → Second chatbot (the "challenger")
├── suite_id      → Which tests were used
├── winner        → "A", "B", or "tie"
├── p_value       → Statistical confidence (lower = more sure)
└── effect_size   → How big was the difference
'''
    para = doc.add_paragraph()
    run = para.add_run(ab_fields)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # AgentChain
    doc.add_heading('AgentChain', level=3)
    doc.add_paragraph('WHAT IT REPRESENTS: A sequence of chatbots that work together.')
    doc.add_paragraph('WHY IT EXISTS: Some systems use multiple chatbots in sequence. For example: Router Bot → Specialist Bot → Summary Bot. We need to test if the whole chain works.')
    chain_fields = '''
AgentChain
├── name          → Name of the chain
├── steps         → Ordered list of chatbots to call
└── fail_fast     → Stop immediately if any step fails?
'''
    para = doc.add_paragraph()
    run = para.add_run(chain_fields)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_page_break()

    # ===========================================
    # Section 6: API Reference
    # ===========================================
    doc.add_heading('6. API Endpoints (How Frontend Talks to Backend)', level=1)

    doc.add_paragraph('WHY WE NEED APIs:')
    doc.add_paragraph(
        'The web browser cannot directly access the database or call chatbots. '
        'It needs to go through the backend server. APIs are the "doors" that allow this communication. '
        'Each API endpoint does one specific thing.'
    )

    doc.add_heading('How to Read This Section', level=2)
    doc.add_paragraph('• GET = Retrieve information (read-only)')
    doc.add_paragraph('• POST = Create something new or trigger an action')
    doc.add_paragraph('• PUT = Update existing information')
    doc.add_paragraph('• DELETE = Remove something')

    doc.add_heading('Health Check', level=2)
    doc.add_paragraph('GET /api/health → Returns {"status": "healthy"} if server is running')
    doc.add_paragraph('WHY: So monitoring tools can check if the service is up.')

    doc.add_heading('Running Tests', level=2)
    test_apis = [
        ('POST /api/test', 'Run a single test against a chatbot', 'The main endpoint - send question, get evaluated response'),
        ('POST /api/batch', 'Run multiple tests at once', 'Faster than calling /api/test many times'),
    ]
    for endpoint, desc, why in test_apis:
        doc.add_paragraph(f'{endpoint}').runs[0].bold = True
        doc.add_paragraph(f'What: {desc}')
        doc.add_paragraph(f'Why: {why}')
        doc.add_paragraph('')

    doc.add_heading('Managing Test Suites', level=2)
    suite_apis = [
        ('GET /api/suites', 'List all test suites', 'Show users what suites exist'),
        ('POST /api/suites', 'Create a new test suite', 'Let users organize tests into groups'),
        ('GET /api/suites/{id}', 'Get one suite with all its tests', 'Load a suite for viewing or running'),
        ('DELETE /api/suites/{id}', 'Delete a test suite', 'Clean up suites no longer needed'),
        ('POST /api/suites/{id}/run', 'Run all tests in a suite', 'Execute an entire test collection at once'),
        ('PUT /api/suites/{id}/tests', 'Replace all tests (JSON upload)', 'Bulk import tests from JSON file'),
    ]
    for endpoint, desc, why in suite_apis:
        doc.add_paragraph(f'{endpoint}').runs[0].bold = True
        doc.add_paragraph(f'What: {desc} | Why: {why}')

    doc.add_heading('Managing Agents (Chatbots)', level=2)
    agent_apis = [
        ('GET /api/agents', 'List registered chatbots', 'Show all chatbots available for testing'),
        ('POST /api/agents', 'Register a new chatbot', 'Add a chatbot to the system'),
        ('DELETE /api/agents/{id}', 'Remove a chatbot', 'Clean up old chatbots'),
        ('POST /api/agents/{id}/test', 'Quick test a chatbot', 'Verify a chatbot is working'),
    ]
    for endpoint, desc, why in agent_apis:
        doc.add_paragraph(f'{endpoint}').runs[0].bold = True
        doc.add_paragraph(f'What: {desc} | Why: {why}')

    doc.add_heading('A/B Testing', level=2)
    ab_apis = [
        ('GET /api/ab-comparisons', 'List past comparisons', 'See history of A/B tests'),
        ('POST /api/ab-comparisons', 'Run a new A/B comparison', 'Compare two chatbots with statistical analysis'),
    ]
    for endpoint, desc, why in ab_apis:
        doc.add_paragraph(f'{endpoint}').runs[0].bold = True
        doc.add_paragraph(f'What: {desc} | Why: {why}')

    doc.add_heading('Chain Testing', level=2)
    chain_apis = [
        ('GET /api/chains', 'List agent chains', 'Show configured chatbot pipelines'),
        ('POST /api/chains', 'Create a new chain', 'Define a sequence of chatbots'),
        ('POST /api/chains/{id}/run', 'Execute a chain', 'Test the entire pipeline'),
    ]
    for endpoint, desc, why in chain_apis:
        doc.add_paragraph(f'{endpoint}').runs[0].bold = True
        doc.add_paragraph(f'What: {desc} | Why: {why}')

    doc.add_heading('History & Analytics', level=2)
    history_apis = [
        ('GET /api/history', 'Get test history (with pagination)', 'Show past results, filterable by date/agent/status'),
        ('GET /api/analytics', 'Get statistics', 'Show summary: total tests, pass rate, trends'),
    ]
    for endpoint, desc, why in history_apis:
        doc.add_paragraph(f'{endpoint}').runs[0].bold = True
        doc.add_paragraph(f'What: {desc} | Why: {why}')

    doc.add_page_break()

    # ===========================================
    # Section 7: Frontend
    # ===========================================
    doc.add_heading('7. The User Interface', level=1)

    doc.add_paragraph('WHERE TO FIND IT:')
    doc.add_paragraph('agent_eval/web/templates/index.html').runs[0].font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('WHY ONE FILE:')
    doc.add_paragraph(
        'The entire UI is a single HTML file. This is intentional. It means: '
        '(1) No build step needed - just edit and refresh. '
        '(2) Easy to deploy - just one file. '
        '(3) Easy to understand - everything in one place.'
    )

    doc.add_heading('The Tabs', level=2)
    doc.add_paragraph(
        'The UI is organized into tabs. Each tab serves a specific purpose:'
    )

    tabs = [
        ('Quick Test',
         'Run a single test right now',
         'Users want to quickly check if their chatbot works without setting up suites'),

        ('Agents',
         'Manage registered chatbots',
         'Users need to add, remove, and test their chatbots'),

        ('Compare',
         'A/B testing and multi-chatbot comparison',
         'Users need to scientifically compare chatbot versions'),

        ('Chains',
         'Build and test chatbot pipelines',
         'Users with orchestrator systems need to test the whole flow'),

        ('Test Suites',
         'Manage test collections',
         'Users want to save tests and run them repeatedly'),

        ('Analytics',
         'View performance dashboards',
         'Users want to see trends and overall health metrics'),

        ('History',
         'Browse past test results',
         'Users need to find and review past evaluations'),
    ]

    tabs_table = doc.add_table(rows=1, cols=3)
    tabs_table.style = 'Table Grid'
    hdr = tabs_table.rows[0].cells
    hdr[0].text = 'Tab'
    hdr[1].text = 'What It Does'
    hdr[2].text = 'Why Users Need It'
    for tab, what, why in tabs:
        row = tabs_table.add_row().cells
        row[0].text = tab
        row[1].text = what
        row[2].text = why

    doc.add_heading('Key JavaScript Functions', level=2)
    doc.add_paragraph('The JavaScript is at the bottom of the HTML file. Key functions:')

    js_funcs = [
        ('openModal() / closeModal()', 'Show and hide popup dialogs'),
        ('loadAgents() / loadSuites() / loadHistory()', 'Fetch data from backend and display it'),
        ('runTest()', 'Execute a test and show results'),
        ('registerAgent()', 'Register a new chatbot'),
        ('createChainFromBuilder()', 'Create a chain from the visual builder'),
    ]
    for func, desc in js_funcs:
        doc.add_paragraph(f'{func} → {desc}')

    doc.add_page_break()

    # ===========================================
    # Section 8: Database
    # ===========================================
    doc.add_heading('8. Database Tables', level=1)

    doc.add_paragraph('DATABASE FILE LOCATION:')
    doc.add_paragraph('~/.agent_eval/data.db').runs[0].font.name = 'Courier New'
    doc.add_paragraph('(In your home folder, inside .agent_eval folder)')

    doc.add_paragraph('')
    doc.add_paragraph('WHY SQLite:')
    doc.add_paragraph(
        'SQLite is a simple database that needs no installation. It\'s just a file. '
        'This makes the tool easy to set up - no database server required.'
    )

    doc.add_heading('Tables', level=2)

    tables = [
        ('suites', 'Test suite definitions', 'name, description, created_at'),
        ('tests', 'Individual test cases', 'suite_id, input, expected, context'),
        ('results', 'Evaluation outcomes', 'input, output, score, passed, latency_ms'),
        ('agents', 'Registered chatbots', 'name, endpoint, auth_type, version'),
        ('ab_comparisons', 'A/B test results', 'agent_a_id, agent_b_id, winner, p_value'),
        ('chains', 'Agent pipelines', 'name, steps (JSON array)'),
        ('chain_runs', 'Chain execution history', 'chain_id, status, passed_tests'),
    ]

    tables_tbl = doc.add_table(rows=1, cols=3)
    tables_tbl.style = 'Table Grid'
    hdr = tables_tbl.rows[0].cells
    hdr[0].text = 'Table'
    hdr[1].text = 'Purpose'
    hdr[2].text = 'Key Columns'
    for name, purpose, cols in tables:
        row = tables_tbl.add_row().cells
        row[0].text = name
        row[1].text = purpose
        row[2].text = cols

    doc.add_page_break()

    # ===========================================
    # Section 9: Getting Started
    # ===========================================
    doc.add_heading('9. Getting Started', level=1)

    doc.add_heading('What You Need', level=2)
    doc.add_paragraph('• Python 3.11 or newer')
    doc.add_paragraph('• Poetry (recommended) or pip')
    doc.add_paragraph('• Optional: OpenAI API key (for better evaluations)')

    doc.add_heading('Installation', level=2)
    install = '''
# 1. Clone the code
git clone <repository-url>
cd agent-eval-mvp

# 2. Install dependencies
poetry install

# 3. (Optional) Enable LLM-powered evaluations
poetry install --extras deepeval
export OPENAI_API_KEY=your-key-here
'''
    para = doc.add_paragraph()
    run = para.add_run(install)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('Running the Server', level=2)
    run_cmd = '''
# Start the server
poetry run uvicorn agent_eval.web.app:app --reload --port 8000

# Open in browser
http://localhost:8000
'''
    para = doc.add_paragraph()
    run = para.add_run(run_cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('Your First Test', level=2)
    steps = [
        'Open http://localhost:8000 in your browser',
        'Go to the "Agents" tab',
        'Click "+ Register Agent" and enter your chatbot\'s URL',
        'Go to the "Quick Test" tab',
        'Type a question and click "Run Test"',
        'See the results with pass/fail status and quality scores',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ===========================================
    # Section 10: Extending
    # ===========================================
    doc.add_heading('10. How to Add New Features', level=1)

    doc.add_heading('Adding a New Evaluation Metric', level=2)
    doc.add_paragraph('WHY: You might want to measure something specific like "response length" or "brand voice consistency".')
    doc.add_paragraph('')
    doc.add_paragraph('STEPS:')
    doc.add_paragraph('1. Open core/evaluator.py')
    doc.add_paragraph('2. Add your metric to the METRICS dictionary')
    doc.add_paragraph('3. Add a handler function _heuristic_your_metric()')
    doc.add_paragraph('4. The metric will automatically appear in the UI')

    doc.add_heading('Adding Support for a New Request Format', level=2)
    doc.add_paragraph('WHY: A new chatbot might expect a format the Executor doesn\'t know.')
    doc.add_paragraph('')
    doc.add_paragraph('STEPS:')
    doc.add_paragraph('1. Open core/executor.py')
    doc.add_paragraph('2. Find the _get_payloads() method')
    doc.add_paragraph('3. Add your new format to the list of payloads to try')

    doc.add_heading('Adding Support for a New File Format', level=2)
    doc.add_paragraph('WHY: Users might want to upload files in formats we don\'t support yet.')
    doc.add_paragraph('')
    doc.add_paragraph('STEPS:')
    doc.add_paragraph('1. Open core/file_parser.py')
    doc.add_paragraph('2. Add the extension to SUPPORTED_EXTENSIONS')
    doc.add_paragraph('3. Add a _parse_yourformat() method')
    doc.add_paragraph('4. Update the parse() method to call it')

    doc.add_heading('Adding a New API Endpoint', level=2)
    doc.add_paragraph('WHY: You might need new functionality accessible via API.')
    doc.add_paragraph('')
    doc.add_paragraph('STEPS:')
    doc.add_paragraph('1. Open web/app.py')
    doc.add_paragraph('2. Create a Pydantic model for request/response')
    doc.add_paragraph('3. Add a function with @app.get() or @app.post() decorator')
    doc.add_paragraph('4. The endpoint is now live')

    # Footer
    doc.add_page_break()
    doc.add_heading('Support & Contact', level=1)
    doc.add_paragraph('For questions or issues, contact the development team.')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('Eli Lilly and Company')
    doc.add_paragraph('Internal Use Only')
    doc.add_paragraph('')
    doc.add_paragraph('Lilly Agent Eval v3.0.0')

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_path = Path(__file__).parent / 'Lilly_Agent_Eval_Developer_Guide.docx'
    doc.save(str(output_path))
    print(f'Document saved to: {output_path}')
