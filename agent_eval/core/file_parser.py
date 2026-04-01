"""
File parser module for extracting text from various document formats.

Supports: PDF, Markdown, DOCX, XLSX, CSV
"""

import io
import csv
from pathlib import Path
from typing import List, Optional
import logging

logger = logging.getLogger(__name__)


class FileParser:
    """
    Parse various file formats and extract text content.

    Supported formats:
    - PDF (.pdf)
    - Markdown (.md)
    - Word Document (.docx)
    - Excel (.xlsx, .xls)
    - CSV (.csv)
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.md', '.docx', '.xlsx', '.xls', '.csv', '.txt'}

    @classmethod
    def parse(cls, file_content: bytes, filename: str) -> str:
        """
        Parse file content and extract text.

        Args:
            file_content: Raw file bytes
            filename: Original filename (used to determine format)

        Returns:
            Extracted text content
        """
        ext = Path(filename).suffix.lower()

        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {cls.SUPPORTED_EXTENSIONS}")

        try:
            if ext == '.pdf':
                return cls._parse_pdf(file_content)
            elif ext == '.md' or ext == '.txt':
                return cls._parse_text(file_content)
            elif ext == '.docx':
                return cls._parse_docx(file_content)
            elif ext in ['.xlsx', '.xls']:
                return cls._parse_excel(file_content)
            elif ext == '.csv':
                return cls._parse_csv(file_content)
            else:
                return cls._parse_text(file_content)
        except Exception as e:
            logger.error(f"Error parsing {filename}: {e}")
            raise ValueError(f"Failed to parse {filename}: {str(e)}")

    @classmethod
    def _parse_pdf(cls, content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")

        reader = PdfReader(io.BytesIO(content))
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num + 1} ---")
                text_parts.append(page_text.strip())

        return "\n\n".join(text_parts)

    @classmethod
    def _parse_text(cls, content: bytes) -> str:
        """Parse plain text or markdown file."""
        # Try UTF-8 first, then fall back to latin-1
        try:
            return content.decode('utf-8').strip()
        except UnicodeDecodeError:
            return content.decode('latin-1').strip()

    @classmethod
    def _parse_docx(cls, content: bytes) -> str:
        """Extract text from Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        doc = Document(io.BytesIO(content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @classmethod
    def _parse_excel(cls, content: bytes) -> str:
        """Extract text from Excel file."""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas and openpyxl are required for Excel parsing. Install with: pip install pandas openpyxl")

        # Read all sheets
        excel_file = pd.ExcelFile(io.BytesIO(content))
        text_parts = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            if not df.empty:
                text_parts.append(f"--- Sheet: {sheet_name} ---")
                # Convert DataFrame to string representation
                text_parts.append(df.to_string(index=False))

        return "\n\n".join(text_parts)

    @classmethod
    def _parse_csv(cls, content: bytes) -> str:
        """Extract text from CSV file."""
        # Decode content
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin-1')

        # Parse CSV
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return ""

        # Format as readable text
        text_parts = []

        # If there's a header row, use it
        if rows:
            header = rows[0]
            text_parts.append(" | ".join(str(h) for h in header))
            text_parts.append("-" * 50)

            for row in rows[1:]:
                text_parts.append(" | ".join(str(cell) for cell in row))

        return "\n".join(text_parts)

    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """Get list of supported file formats."""
        return [
            {"extension": ".pdf", "description": "PDF Document"},
            {"extension": ".md", "description": "Markdown File"},
            {"extension": ".txt", "description": "Plain Text"},
            {"extension": ".docx", "description": "Word Document"},
            {"extension": ".xlsx", "description": "Excel Spreadsheet"},
            {"extension": ".csv", "description": "CSV File"},
        ]


def parse_file(file_content: bytes, filename: str) -> str:
    """
    Convenience function to parse a file.

    Args:
        file_content: Raw file bytes
        filename: Original filename

    Returns:
        Extracted text content
    """
    return FileParser.parse(file_content, filename)
